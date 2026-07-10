#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Assemble the full thesis .docx with UITM Rzeszow formatting."""
import os
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

VLASH = "/sessions/eloquent-inspiring-pascal/mnt/vlash"
SIMR = os.path.join(VLASH, "simulation", "results")
FIGS = os.path.join(VLASH, "figures")
OUTDOC = os.path.join(VLASH, "Thesis_VLASH_MacDuyKhanh.docx")

FONT = "Times New Roman"
doc = Document()

# ---------- base style ----------
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
pf = normal.paragraph_format
pf.line_spacing = 1.15
pf.space_after = Pt(6)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def style_heading(name, size):
    st = doc.styles[name]
    st.font.name = FONT
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor(0, 0, 0)
    st.paragraph_format.space_before = Pt(12)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.15
    st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
style_heading("Heading 1", 14)
style_heading("Heading 2", 12)
style_heading("Heading 3", 12)

# ---------- page setup (A4, 2 cm margins) ----------
sec = doc.sections[0]
sec.page_width = Mm(210); sec.page_height = Mm(297)
sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
sec.left_margin = Cm(2); sec.right_margin = Cm(2)
sec.different_first_page_header_footer = True  # no number on title page

# footer page number (centered) on non-first pages
def add_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    run.font.name = FONT; run.font.size = Pt(12)
add_page_number(sec)

# ---------- helpers ----------
def para(text, align="justify", bold=False, italic=False, size=12, after=6, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(after)
    amap = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT}
    p.alignment = amap[align]
    if indent and align == "justify":
        p.paragraph_format.first_line_indent = Cm(0.5)
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    return p

def h1(text):
    p = doc.add_paragraph(text, style="Heading 1")
    p.paragraph_format.page_break_before = True
    return p
def h2(text): return doc.add_paragraph(text, style="Heading 2")
def h3(text): return doc.add_paragraph(text, style="Heading 3")

def equation(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.name = FONT; r.font.size = Pt(12); r.italic = True
    return p

def figure(path, caption, source="own work", width_cm=14.5):
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER  # caption ABOVE object
    cap.paragraph_format.space_before = Pt(10); cap.paragraph_format.space_after = Pt(2)
    rc = cap.add_run(caption); rc.font.name = FONT; rc.font.size = Pt(10)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(); r.add_picture(path, width=Cm(width_cm))
    src = doc.add_paragraph(); src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = src.add_run(f"Source: {source}"); rs.font.name = FONT; rs.font.size = Pt(10)
    src.paragraph_format.space_after = Pt(10)
    return p

def table_caption(text):
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(text); r.font.name = FONT; r.font.size = Pt(10)
    cap.paragraph_format.space_before = Pt(8); cap.paragraph_format.space_after = Pt(2)

def make_table(rows, source="own work", header=True, font_size=10):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            pp = cells[ci].paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp.paragraph_format.space_after = Pt(0); pp.paragraph_format.line_spacing = 1.0
            rr = pp.add_run(str(val)); rr.font.name = FONT; rr.font.size = Pt(font_size)
            rr.bold = (ri == 0 and header)
    src = doc.add_paragraph(); src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = src.add_run(f"Source: {source}"); rs.font.name = FONT; rs.font.size = Pt(10)
    src.paragraph_format.space_after = Pt(10)
    return t

def field(p, instr):
    run = p.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = instr
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Right-click → Update Field"
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(i); run._r.append(sep); run._r.append(t); run._r.append(e)

# ===================================================================
# TITLE PAGE
# ===================================================================
def tp(text, size, bold=False, italic=False, after=6, before=0):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text); r.font.name = FONT; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    return p

tp("UNIVERSITY OF INFORMATION TECHNOLOGY AND MANAGEMENT IN RZESZÓW", 13, bold=True, before=6)
tp("FACULTY OF INFORMATION TECHNOLOGY", 16, bold=True, after=18)
tp("Field of Study: INFORMATION TECHNOLOGY", 12, after=2)
tp("Specialty: Programming", 12, after=40)
tp("Duy Khanh Mac", 14, after=2)
tp("No. of student's record book 72850", 12, after=40)
tp("Real-Time Vision-Language-Action (VLA) Inference for Reactive Robotics", 20, bold=True, italic=True, after=40)
tp("Supervisor: prof. WSIIZ dr inż. Atsuo Murata", 14, after=46)
tp("DIPLOMA THESIS AT FIRST-CYCLE STUDIES", 18, bold=True, after=46)
tp("Rzeszów 2026", 14, bold=True)

# ===================================================================
# TABLE OF CONTENTS (static, page-numbered via 2-pass build)
# ===================================================================
import json
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
PAGEMAP = {}
_pm = os.path.join(VLASH, "pagemap.json")
if os.path.exists(_pm):
    PAGEMAP = json.load(open(_pm))

TOC_ENTRIES = [
    ("Introduction", 1),
    ("CHAPTER 1: Background and Related Work", 1),
    ("1.0 Introduction to the Chapter", 2), ("1.1 Robotics and Real-Time Control", 2),
    ("1.2 Vision-Language Models and VLAs", 2), ("1.3 Action Chunking and Inference Delay", 2),
    ("1.4 Asynchronous Inference in Robotics", 2), ("1.5 Limitations of Existing Approaches", 2),
    ("1.6 Summary of the Chapter", 2),
    ("CHAPTER 2: System Design and Methodology", 1),
    ("2.0 Introduction to the Chapter", 2), ("2.1 Overview and Objectives of the System", 2),
    ("2.2 The VLASH Method Under Study", 2), ("2.3 Design of the Simulation Testbed", 2),
    ("2.4 Implementation", 2), ("2.5 Relation to the Full-Scale Reference Implementation", 2),
    ("2.6 Summary of the Chapter", 2),
    ("CHAPTER 3: Experiments and Results", 1),
    ("3.0 Introduction to the Chapter", 2), ("3.1 Experimental Setup", 2),
    ("3.2 Evaluation Metrics", 2), ("3.3 Simulation Study Results", 2),
    ("3.4 Corroboration with Large-Scale VLASH Results", 2), ("3.5 Discussion", 2),
    ("3.6 Summary of the Chapter", 2),
    ("CHAPTER 4: Conclusion and Future Work", 1),
    ("4.0 Conclusion", 2), ("4.1 Future Work", 2),
    ("Bibliography", 1), ("List of Figures", 1), ("List of Tables", 1), ("Summary", 1),
]
ptoc = doc.add_paragraph(); ptoc.paragraph_format.page_break_before = True
ptoc.paragraph_format.space_after = Pt(8)
rt = ptoc.add_run("Table of Contents"); rt.font.name = FONT; rt.font.size = Pt(14); rt.bold = True
for title, lvl in TOC_ENTRIES:
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(2)
    if lvl == 2:
        p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(17.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(title); r.font.name = FONT; r.font.size = Pt(12); r.bold = (lvl == 1)
    pg = str(PAGEMAP.get(title, "")) if PAGEMAP else ""
    r2 = p.add_run("\t" + pg); r2.font.name = FONT; r2.font.size = Pt(12)

def bullet(text, num=False):
    p = doc.add_paragraph(style=("List Number" if num else "List Bullet"))
    p.paragraph_format.line_spacing = 1.15; p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.font.name = FONT; r.font.size = Pt(12)
    return p

# ===================================================================
# INTRODUCTION
# ===================================================================
h1("Introduction")
para("In recent years, artificial intelligence has been widely adopted in robotics. Modern robots can observe their surroundings and choose appropriate actions. In practice, however, AI-powered robots still experience delays in decision-making — a critical drawback for tasks that demand rapid response — because the robot must wait for its reasoning model to finish before it can act. This produces a “stop–run–stop” behaviour and less than smooth operation. The central question of this thesis is how to enable a robot to think and move at the same time while preserving accuracy, which motivates a new inference framework for Vision-Language-Action (VLA) models.")
para("A VLA model is a class of artificial intelligence model built around three components:", indent=False)
bullet("Vision — images or sensor data from the environment;")
bullet("Language — task descriptions in natural language;")
bullet("Action — the commands executed by the robot.")
para("VLA models let robots “understand” their surroundings through images and human requests through language, then translate that understanding into actions. They typically operate with an action-chunking mechanism: at each step the model generates a short sequence of future actions, the robot executes them, and the model then produces a new sequence and repeats the process.")
para("Under synchronous inference the robot must wait for the model to finish reasoning before continuing, so it is often immobile during that time. Asynchronous inference removes this stall by letting the robot execute previous actions while the model reasons about new ones, allowing near-continuous motion. Naive asynchronous inference, however, has a significant problem: a misalignment between the state used for prediction and the actual state at execution time. While the model performs inference, both the robot and the environment continue to evolve, so the newly predicted actions may no longer suit the current state.")
para("This thesis focuses on future-state-aware asynchronous inference — estimating the future state of the robot at the moment new actions will be executed — as embodied by the VLASH framework [1].")
para("The objectives of this thesis are:", indent=False)
bullet("to present an overview of VLA models and real-time inference mechanisms;")
bullet("to analyse the problem of prediction–execution misalignment in asynchronous inference;")
bullet("to explain the future-state-aware method;")
bullet("to design a simulation system that illustrates and compares three inference mechanisms — synchronous, naive asynchronous, and future-state-aware;")
bullet("to evaluate and discuss the results obtained from the simulation and to corroborate them against the large-scale results reported for VLASH.")
para("The thesis is organised into four chapters. Chapter 1 presents the background and related work on VLA models and real-time robotic inference. Chapter 2 describes the methodology and the design of the simulation system together with the future-state-aware method. Chapter 3 reports the experimental setup, the results of the simulation, and a comparison against the published large-scale results. Chapter 4 concludes the thesis, summarises its contributions, and suggests directions for future work.")

# ===================================================================
# CHAPTER 1
# ===================================================================
h1("CHAPTER 1: Background and Related Work")
h2("1.0 Introduction to the Chapter")
para("The deployment of machine-learning models in robotics presents challenges that differ substantially from conventional computer-vision or natural-language-processing applications. Unlike offline prediction tasks, robotic control demands continuous, real-time interaction with a physical or simulated environment, where delays in computation translate directly into degraded task performance. This chapter provides the theoretical and empirical background needed to understand these challenges and the state of the art in addressing them.")
para("The chapter is organised as follows. Section 1.1 introduces fundamental concepts of real-time robotic control, including control loops, control frequencies, and the action-chunking paradigm. Section 1.2 surveys Vision-Language Models and their extension to robot action generation through VLA models, with attention to the π0.5 model used in this work. Section 1.3 formalises the problem of inference latency. Section 1.4 reviews existing asynchronous inference strategies. Section 1.5 identifies the key limitations of prior approaches that motivate the contributions of this thesis.")

h2("1.1 Robotics and Real-Time Control")
para("Modern robotic systems rely on a continuous feedback loop in which the robot repeatedly senses its environment, computes a response, and executes the resulting actions. This control loop must operate at a frequency sufficient to react to changes before they render prior decisions obsolete. Depending on the task, control frequencies typically range from about 10 Hz for manipulation requiring moderate precision to over 50 Hz for highly dynamic scenarios [7, 8].")
para("Traditional controllers were hand-designed using model-based techniques, where an explicit mathematical model of the robot and environment was used to derive optimal actions. Such approaches require detailed domain knowledge and struggle to generalise across diverse, unstructured tasks. The emergence of deep learning shifted the paradigm toward learned policies, where a neural network πθ maps an observation o_t and robot state s_t to a sequence of actions:")
equation("πθ ( A_t | o_t , s_t )")
para("Here A_t is the predicted action sequence, o_t the current sensory observation (e.g. RGB images), and s_t the proprioceptive state such as joint positions and gripper configuration.")
para("A key development is action chunking [8], in which the policy does not predict a single action but generates a chunk of H consecutive actions in one forward pass:")
equation("A_t = { a_t , a_{t+1} , … , a_{t+H−1} }")
para("where H is the prediction horizon. Of this chunk, only the first K actions — the execution horizon — are sent to the actuators before a new chunk is requested. Action chunking reduces the frequency of model queries and has been shown to improve task completion by smoothing temporal inconsistencies in action prediction [8].")
para("Action chunking nevertheless introduces a bottleneck in synchronous deployment: before the next chunk can execute, the model must complete a full forward pass over the current observation. On modern hardware this takes tens to hundreds of milliseconds — an inference latency of Δ control steps — during which the robot is effectively frozen. This action-stall problem worsens as model size grows. On an NVIDIA RTX 4090, inference for π0.5 takes about 103 ms, more than two control steps at 20 Hz [1].")
para("The stall not only reduces the robot’s effective speed but also degrades its ability to react to sudden environmental changes, since new observations cannot be acted upon until inference completes. Addressing this limitation is the central motivation of this thesis.")

h2("1.2 Vision-Language Models and VLAs")
para("The rapid progress of large language models has created a foundation for multimodal learning. Vision-Language Models (VLMs) extend language models with a visual encoder — typically a Transformer-based image encoder such as SigLIP [14] — whose output tokens are projected into the same embedding space as text tokens and processed jointly by the language backbone. Trained on large-scale image-text pairs, VLMs acquire broad visual grounding. Representative models include PaLI, LLaVA, and PaliGemma [13].")
para("While VLMs excel at perception and reasoning, they lack a mechanism to produce low-level robot actions. Vision-Language-Action models (VLAs) bridge this gap by augmenting a VLM backbone with an action expert that generates sequences of continuous control commands. The VLM processes the current image(s) and the natural-language task description, producing a contextual representation; the action expert conditions on this representation together with the robot state to output an action chunk. This enables end-to-end training from human demonstrations.")
para("Several VLA models have been proposed. RT-2 [12] was among the first to show that a web-pretrained VLM could be fine-tuned for manipulation. π0 [2] introduced a flow-matching action expert with high dexterity on contact-rich tasks. GR00T N1 [15] scaled the paradigm to humanoid robots. The model used in this work, π0.5 [3], is a refinement designed for instruction-following across diverse manipulation scenarios.")
para("Figure 1 summarises the π0.5 architecture. The vision-language backbone is PaliGemma, which encodes up to two camera images with a SigLIP encoder and processes the patch tokens alongside tokenised instructions. The action expert is a separate Gemma model that generates action tokens. The two are connected by joint attention [2]: at each layer the action expert’s queries attend over its own key-value pairs and those of the backbone, letting the action expert read the full visual and linguistic context while keeping inference efficient.")
figure(os.path.join(FIGS, "fig_concept_pi05.png"),
       "Fig. 1. Architecture of the π0.5 vision-language-action model: a PaliGemma backbone and a Gemma action expert coupled by joint attention, with state conditioning and a flow-matching action head.",
       source="own work, based on [1, 3]")
para("Robot-state conditioning is achieved through AdaRMS, an adaptive root-mean-square normalisation whose scale and shift parameters are modulated by a linear projection of the robot state and a diffusion timestep embedding. Unlike earlier VLAs that discretised actions into tokens, π0.5 generates continuous actions through flow matching [11], learning a vector field vθ(x, t) that defines a probability-flow ODE:")
equation("dx / dt = vθ ( x , t )")
para("Integrating this ODE over t ∈ [0, 1] maps a noise sample to a predicted action chunk; flow matching admits straighter paths than diffusion, enabling accurate generation in as few as ten integration steps, which is critical for low-latency control. In practice the pretrained π0.5 weights are adapted to a target task through parameter-efficient fine-tuning with Low-Rank Adaptation (LoRA) [10], updating fewer than 1% of parameters.")

h2("1.3 Action Chunking and Inference Delay")
para("As introduced in Section 1.1, action-chunking policies generate a block of H actions in a single query. To reason precisely about timing it is useful to define two intervals over discrete control steps. At decision time t, the policy observes o_t and s_t and produces a chunk spanning the prediction interval I_pred = [t, t + K), where K ≤ H is the execution horizon. In a purely synchronous system the forward pass begins at t and completes after Δ control steps, during which the robot is halted; the first K actions then form the execution interval I_exec = [t + Δ, t + Δ + K).")
para("This reveals the core prediction–execution misalignment: the policy conditioned its output on the state observed at t, but the actions do not begin executing until t + Δ, by which point the robot state has evolved. The severity of the misalignment grows with Δ. Table 1 summarises measured inference latencies for π0.5 on representative hardware [1].")
table_caption("Tab. 1. Measured inference latency of π0.5 on representative hardware platforms.")
make_table([
    ["Hardware", "Inference time", "Control steps (at 10 Hz)"],
    ["NVIDIA RTX 4090", "~103 ms", "~1 step"],
    ["NVIDIA RTX 5090", "~30 ms", "< 1 step"],
    ["Consumer laptop GPU", "~200–400 ms", "2–4 steps"],
], source="based on [1]")
para("At a 10 Hz control frequency an RTX 4090 incurs a delay of roughly one step — already meaningful for reactive tasks — while a laptop-class GPU can reach three to four steps. The combination of synchronous stalling and misalignment produces two failure modes: action stall, which reduces the effective control frequency and causes visible hesitation, and miscalibrated actions, where the policy issues commands for a configuration the robot no longer occupies, leading to compounding errors in contact-rich manipulation. Larger models achieve higher accuracy but amplify both effects; resolving this tension is the central problem of this thesis.")

h2("1.4 Asynchronous Inference in Robotics")
para("The stall has a conceptually simple remedy: decouple the inference thread from the execution thread so that the model computes the next chunk while the robot executes the current one. Under this asynchronous paradigm, if inference finishes before the current chunk is exhausted, the transition incurs no idle time and the robot moves continuously. This can reduce the effective reaction latency from a full Δ-step wait to the time remaining in the current chunk [1]. Figure 2 contrasts the three inference paradigms studied in this thesis.")
figure(os.path.join(FIGS, "fig_concept_timeline.png"),
       "Fig. 2. Three inference paradigms. Synchronous inference stalls the robot; naive asynchronous inference removes the stall but leaves prediction–execution misalignment; VLASH conditions on the rolled-forward future state, removing the misalignment.",
       source="own work, based on [1]")
para("Naive asynchronous inference, as implemented in SmolVLA [4], queries the policy with the observation and state captured when inference begins and swaps in the resulting chunk as soon as it is ready, with no adjustment for elapsed time. It eliminates stall and achieves low latency but does not address misalignment: if inference begins at t and finishes at t + Δ, the policy conditioned on s_t while the actions execute from s_{t+Δ}. The gap s_{t+Δ} − s_t grows with Δ, producing jittery, laggy motion in contact-rich tasks.")
para("Real-Time Chunking (RTC) [5] addresses misalignment by modifying the chunk at runtime: actions already executed are frozen and the remaining suffix is regenerated (inpainted) conditioned on the executed prefix. While RTC reduces misalignment, the inpainting step requires an additional forward pass at the moment the chunk is consumed, adding latency precisely when fast response matters most, and it requires the policy to support conditional inpainting.")
para("A2C2 [6] takes a different route: a lightweight correction module is added on top of the base policy and trained to predict per-step residual corrections to the pending chunk from the latest observation. It interfaces with an off-the-shelf VLA without retraining the base model, but it still adds a trained module and per-step runtime computation, and it does not explicitly model the robot’s future kinematic state. In summary, naive async leaves misalignment unaddressed, RTC corrects it at the cost of runtime inpainting, and A2C2 adds a correction head and per-step overhead; none simultaneously achieves zero overhead, no added module, and explicit future-state compensation.")

h2("1.5 Limitations of Existing Approaches")
para("The strategies above each advance over synchronous deployment yet each carries a distinct limitation. Naive asynchronous inference trades the stall for an unresolved misalignment that compounds as Δ grows, performing worse than synchronous inference on precise or dynamic tasks despite its lower latency [1]. RTC corrects misalignment but adds inpainting overhead on the critical path and requires inpainting-capable training. A2C2 reduces error during execution but adds a trained correction module and per-step computation, and it operates on the action representation rather than modelling the future kinematic state. A practical constraint shared by all asynchronous approaches is the need for multi-threaded execution.")
para("Taken together, these limitations define a clear gap. An ideal solution should satisfy four properties simultaneously: (1) zero runtime overhead beyond the standard forward pass; (2) no architectural modification, so that any pretrained checkpoint can be used; (3) explicit misalignment compensation, conditioning the policy on the future state at which its predictions will execute; and (4) plug-and-play deployment as a fine-tuning procedure. No existing method satisfies all four. This gap motivates the approach studied in Chapter 2.")

h2("1.6 Summary of the Chapter")
para("This chapter established the foundations for the thesis. Real-time control imposes strict latency constraints; the action-chunking paradigm reduces query frequency but merely defers the latency problem. Modern VLAs such as π0.5 combine powerful vision-language backbones with continuous action generation via flow matching, but their inference time on practical hardware creates a gap of one to four control steps between prediction and execution. This gap — the prediction–execution misalignment between I_pred and I_exec — is the root cause of degraded accuracy under asynchronous deployment, not the stalling alone. Existing strategies each address only part of the problem. Chapter 2 presents the methodology and the simulation system used to study a technique — state rollforward — that conditions the policy on an analytically computed future state at no additional inference cost.")

# ===================================================================
# CHAPTER 2
# ===================================================================
h1("CHAPTER 2: System Design and Methodology")
h2("2.0 Introduction to the Chapter")
para("Chapter 1 established why synchronous VLA inference is insufficient for reactive robotics and why existing asynchronous approaches fail to achieve, at once, zero runtime overhead, architectural compatibility, and accurate misalignment compensation. This chapter describes the methodology and the system designed in this thesis to study and validate a solution.")
para("The contribution is twofold. First, the thesis studies the VLASH method [1] and presents it in a unified form (Section 2.2). Second — and central to the practical part — it designs and implements an original simulation testbed (Sections 2.3–2.4) that reproduces the prediction–execution misalignment phenomenon in a controlled, fully reproducible setting and quantifies how synchronous, naive asynchronous, and future-state-aware inference behave as the inference delay grows. The testbed deliberately abstracts away the heavy VLA model so that the misalignment mechanism can be isolated and measured without large GPU clusters.")
para("Section 2.1 states the goals and scope. Section 2.2 summarises the VLASH method. Section 2.3 describes the testbed: the task, the state and action representation, the inference-delay model, the three controllers, and the metrics. Section 2.4 documents the implementation and tools. Section 2.5 relates the testbed to the full-scale reference implementation and the author’s hardware. Section 2.6 summarises the chapter.")

h2("2.1 Overview and Objectives of the System")
para("The system developed in this thesis is a lightweight, CPU-only simulation testbed whose purpose is to demonstrate, measure, and compare the three inference strategies of Chapter 1 under a common, controlled task. The guiding question is operational: as the inference delay Δ increases, how does each strategy degrade in accuracy, in motion smoothness, and in the underlying prediction–execution misalignment?")
para("To isolate the mechanism, the testbed makes three simplifying assumptions, each justified in Chapter 1. The robot is modelled as a point end-effector moving in a 2-D plane, which preserves the delta-action property that makes state rollforward exact while removing the cost of a full kinematic or contact simulator. The policy is replaced by a deterministic open-loop chunk planner, which retains the essential structure of action chunking — a block of actions committed in advance — that is the structural cause of misalignment. Finally, the inference delay Δ is modelled directly as an integer number of control steps so that it can be swept as an independent variable. Under these assumptions the testbed is not a substitute for the full-scale experiments but an instrument that illustrates and validates the mechanism; the large-scale results are used for corroboration in Section 3.4.")

h2("2.2 The VLASH Method Under Study")
h3("2.2.1 State rollforward")
para("Under asynchronous inference, when the policy is queried at step t the robot is still executing the pending actions of the current chunk. Because actions are delta joint positions (s_{t+1} = s_t + a_t), the future state at which the new chunk will begin executing can be computed exactly and at negligible cost as a vector sum over the known pending actions:")
equation("s_{t+Δ} = s_t + Σ_{i=0}^{Δ−1} a_{t+i}")
para("The policy is then conditioned on this rolled-forward state instead of the stale current state, πθ ( A_{t+Δ} | o_t , s_{t+Δ} ). This single substitution is the entire runtime change; it removes the state-misalignment error without any additional model call. Figure 3 illustrates the computation.")
figure(os.path.join(FIGS, "fig_concept_rollforward.png"),
       "Fig. 3. State rollforward. The pending delta actions are already known, so the execution-time state is obtained exactly by a single vector summation, at zero inference overhead.",
       source="own work, based on [1]", width_cm=14.0)
h3("2.2.2 Temporal-offset augmentation")
para("State rollforward only helps if the policy has learned to exploit a future state that is temporally offset from the observation. VLASH achieves this during fine-tuning by sampling a random offset δ ~ Uniform{0, …, max_delay_steps} for each training sample and forming the pair (o_t, s_{t+δ}) → a_{t+δ : t+δ+H}. Training across the full range of offsets makes the policy robust to any inference delay, including the synchronous case δ = 0.")
h3("2.2.3 Shared-observation optimisation")
para("Because every offset for a given timestep shares the same ~700-token observation, VLASH packs N offset branches into a single forward pass with a block-sparse attention mask: observation tokens are encoded once and attended to by all branches, while each state–action branch is isolated from the others. This yields a reported 3.26× training speed-up at equal effective batch size.")
h3("2.2.4 Action quantization")
para("Grouping q consecutive delta actions into a macro-action â_i = a_{iq} + … + a_{(i+1)q−1} lets the policy emit H/q outputs per chunk, trading temporal resolution for throughput. The VLASH authors report a 2.03× end-to-end speed-up at q = 2 with negligible accuracy loss [1]. Of these four components, the testbed focuses on state rollforward, the component that governs deployment-time accuracy; temporal-offset augmentation is modelled implicitly through a planner that can use whichever conditioning state it is given.")

h2("2.3 Design of the Simulation Testbed")
h3("2.3.1 The tracking task")
para("The testbed implements a reactive reaching task. A point end-effector with position s ∈ ℝ² must keep itself on a moving target g_t ∈ ℝ². The target follows a staircase profile: it holds a fixed position for a randomised interval of about 40 control steps and then jumps by a random displacement of magnitude up to 0.25 before holding again. The held segments turn any overshoot caused by misalignment into a pure tracking error — a smoothly drifting target was rejected during design because it lets naive overshoot accidentally anticipate the target, masking the effect under study — while the jumps probe how quickly each strategy reacts.")
h3("2.3.2 State and action representation")
para("Consistent with the delta-action property exploited by VLASH, the action is an incremental displacement a ∈ ℝ², clipped to a maximum magnitude per step, and the dynamics are the exact summation s_{t+1} = s_t + a_t. This is the minimal representation for which state rollforward is exact.")
h3("2.3.3 The chunk planner")
para("Given a conditioning state s_cond and an observed target g_obs, the planner produces an open-loop chunk of K identical delta actions that, if executed from s_cond, drive the end-effector onto g_obs:")
equation("a_k = clip ( (g_obs − s_cond) / K , −MAX_STEP , +MAX_STEP ) ,  k = 0 … K−1")
para("The chunk is open-loop: the same fixed deltas are applied wherever the robot actually is. If the chunk is later executed from a different start state, every delta is unchanged but the landing point shifts by exactly the misalignment ε = s_exec_start − s_cond, which makes the planner an ideal probe for the mechanism.")
h3("2.3.4 The three controllers")
para("All three controllers share the planner; they differ only in when inference runs and which state it conditions on. In synchronous inference (Algorithm 1) the robot freezes for Δ steps while inference runs, then executes a chunk planned from the state it just observed: there is no misalignment but the robot is idle and reacts slowly. In naive asynchronous inference (Algorithm 2) the robot never freezes; inference for the next chunk is launched Δ steps before the current chunk is exhausted and conditioned on the stale current state s_t, so the incurred misalignment is ε = ‖s_{t+Δ} − s_t‖. In VLASH (Algorithm 3) the conditioning state is rolled forward through the known pending deltas, so the rolled-forward state equals the true execution-start state and the modelled misalignment is ε = 0 by construction. The target observation g_obs remains stale in all asynchronous cases — the simulation cannot see the future scene — faithfully reproducing VLASH’s stated limitation that rollforward compensates for robot state but not for visual change.")
h3("2.3.5 Evaluation metrics")
para("Four metrics are recorded per episode and averaged over trials. Misalignment ε is the mean ‖s_exec_start − s_assumed‖ per chunk swap; it directly measures prediction–execution misalignment. Tracking error is the mean Euclidean distance ‖s − g‖ over the episode. Success rate is the fraction of control steps for which ‖s − g‖ lies within a fixed tolerance. Smoothness (jerk) is the mean magnitude of change between consecutive actions, with smaller values indicating smoother motion. Reaction latency is treated analytically in this thesis rather than measured from the testbed, because the kinematic abstraction under-represents the true synchronous stall and would therefore understate the asynchronous reaction advantage established in Chapter 1.")

h2("2.4 Implementation")
para("The testbed is implemented in Python 3 using NumPy for the dynamics and metrics and Matplotlib for the figures; no machine-learning framework or GPU is required, so the experiments run in seconds on a standard laptop. The implementation is a single documented module organised into the target generator, the chunk planner, the three controllers, the metric functions, and an experiment driver that sweeps the delay and aggregates results.")
para("The principal parameters are the execution horizon K = 8, the action clip MAX_STEP = 0.05, the target jump interval (about 40 steps) and magnitude (0.25), the success tolerance (0.03), and the episode length (600 steps). Each delay level Δ ∈ {0, 1, 2, 3, 4} is evaluated over 50 independent trials with distinct random seeds, and every reported value is the mean over those trials. All randomness derives from a single base seed, so every number can be regenerated exactly by re-running the module; raw per-trial results and aggregated summaries are written to CSV alongside the figures. The tools used are Python 3.11 with NumPy, pandas, and Matplotlib, developed under Git version control in Visual Studio Code.")

h2("2.5 Relation to the Full-Scale Reference Implementation")
para("The abstracted testbed is intentionally a small model of a larger system. The full-scale reference — used for the corroborating results in Section 3.4 — is the VLASH codebase [1] built on the HuggingFace LeRobot framework [19], which fine-tunes the π0.5 model [3] (a PaliGemma backbone [13] with a Gemma action expert and a flow-matching action head [11]) on the LIBERO benchmark [7]. There the state is the eight-dimensional Franka Panda configuration, the observation is two 224×224 RGB images plus a language instruction, and the action is an eight-dimensional delta joint command — the same delta property the testbed relies on.")
para("This thesis also targets eventual execution on the author’s own hardware, an NVIDIA RTX 3050 (6 GB) laptop, for which the reference configuration uses 4-bit QLoRA quantization and LoRA adapters to fit the model in memory. A full fine-tune and LIBERO evaluation on this hardware is identified as the natural next step in Section 4.1; the simulation testbed is the feasible, self-contained practical contribution delivered within the scope of this thesis.")

h2("2.6 Summary of the Chapter")
para("This chapter presented the methodology. It summarised the VLASH method — state rollforward, temporal-offset augmentation, the shared-observation optimisation, and action quantization — and described the original contribution: a lightweight, reproducible simulation testbed implementing a reactive reaching task with delta actions and an explicit inference-delay model, realising three controllers that differ only in the state on which they condition. Four metrics were defined to quantify behaviour as the delay grows. Chapter 3 reports the results of this testbed and corroborates them against the large-scale VLASH evaluation.")

# ===================================================================
# CHAPTER 3
# ===================================================================
h1("CHAPTER 3: Experiments and Results")
h2("3.0 Introduction to the Chapter")
para("Chapters 1 and 2 established the motivation for future-state-aware asynchronous inference and the simulation testbed designed to study it. This chapter reports the experimental results in two complementary parts. The first and primary part (Sections 3.1–3.3) presents the results of the original simulation testbed, which isolates and quantifies the prediction–execution misalignment mechanism under controlled conditions. The second part (Section 3.4) corroborates those findings against the large-scale evaluation reported for VLASH [1] on the LIBERO and Kinetix benchmarks and on real robots. Section 3.5 discusses the findings and limitations; Section 3.6 summarises the chapter.")

h2("3.1 Experimental Setup")
para("Simulation study. All results in Sections 3.2–3.3 are produced by the testbed of Chapter 2. The inference delay is swept over Δ ∈ {0, 1, 2, 3, 4} control steps; for each delay, every method is evaluated over 50 independent episodes of 600 control steps with a staircase target that jumps at randomised intervals of roughly 40 steps. The execution horizon is K = 8, the per-step action magnitude is capped at 0.05, and the success tolerance is 0.03. All reported values are means over the 50 trials and are exactly reproducible from a single random seed.")
para("Large-scale reference. The corroborating results in Section 3.4 are those reported for π0.5 [1, 3]. The LIBERO experiments fine-tune π0.5 for 30K iterations at batch size 32 with execution horizon K = 5; latency is measured on a laptop RTX 4090 at 103 ms per forward pass over two images. The real-world experiments use H = 50, K = 24 at 30 Hz on Galaxea R1 Lite and LeRobot SO-101 arms; the reaction-speed measurements use K = 25 at 50 Hz across RTX 5090, 4090, and 5070 GPUs.")

h2("3.2 Evaluation Metrics")
para("The simulation study reports four metrics, defined in Section 2.3.5 and restated here. Misalignment ε is the mean gap between the state a chunk was conditioned on and the state from which it actually begins executing. Tracking error is the mean distance between the end-effector and the target. Success rate is the fraction of control steps within tolerance of the target. Smoothness is the mean magnitude of change between consecutive actions, smaller being smoother. The large-scale reference results in Section 3.4 additionally use task success rate, completion time, speed-up, and reaction latency as defined by the VLASH authors [1].")

h2("3.3 Simulation Study Results")
h3("3.3.1 Misalignment grows with delay but is eliminated by rollforward")
para("Figure 4 plots the measured misalignment ε against the inference delay. For naive asynchronous inference, ε grows monotonically with Δ — from 0.005 at Δ = 1 to 0.036 at Δ = 4 — because the chunk is conditioned on a state the robot has already left by the time the chunk executes. For both synchronous inference and VLASH, ε is identically zero at every delay: synchronous inference never executes from a different state than it planned for, and VLASH rolls the state forward so that the conditioning state matches the execution-start state exactly. This is the clearest confirmation of the central hypothesis. Table 2 reports all four metrics.")
figure(os.path.join(SIMR, "fig_misalignment.png"),
       "Fig. 4. Prediction–execution misalignment ε versus inference delay. Naive async grows with Δ; VLASH and synchronous inference remain at zero.")
table_caption("Tab. 2. Simulation results (mean over 50 trials per delay).")
make_table([
    ["Δ", "Method", "Misalign. ε", "Track. error", "Success (%)", "Jerk (×10⁻³)"],
    ["0", "Sync", "0.000", "0.029", "78.0", "1.02"],
    ["0", "Naive-Async", "0.000", "0.033", "75.8", "1.02"],
    ["0", "VLASH", "0.000", "0.033", "75.8", "1.02"],
    ["1", "Sync", "0.000", "0.032", "76.9", "1.03"],
    ["1", "Naive-Async", "0.005", "0.038", "74.9", "1.14"],
    ["1", "VLASH", "0.000", "0.034", "75.8", "1.03"],
    ["2", "Sync", "0.000", "0.034", "75.8", "1.03"],
    ["2", "Naive-Async", "0.013", "0.049", "60.8", "1.21"],
    ["2", "VLASH", "0.000", "0.038", "73.4", "1.02"],
    ["3", "Sync", "0.000", "0.037", "74.7", "1.07"],
    ["3", "Naive-Async", "0.023", "0.061", "47.5", "1.28"],
    ["3", "VLASH", "0.000", "0.043", "71.0", "1.06"],
    ["4", "Sync", "0.000", "0.038", "73.8", "1.03"],
    ["4", "Naive-Async", "0.036", "0.071", "35.4", "1.36"],
    ["4", "VLASH", "0.000", "0.046", "69.5", "1.02"],
], font_size=9)
h3("3.3.2 Task accuracy")
para("Figures 5 and 6 show success rate and tracking error against delay. Synchronous inference is the accuracy upper bound, declining only marginally from 78.0% to 73.8% across the delay range (its small loss reflects slower reaction to target jumps, not misalignment). Naive asynchronous inference starts comparably but collapses as the delay grows, falling to 35.4% success at Δ = 4 as accumulated misalignment pushes the open-loop chunk past the target. VLASH tracks the synchronous curve closely, retaining 69.5% success at Δ = 4 — roughly double the naive baseline — while staying asynchronous. The tracking-error figure tells the same story: at Δ = 4 the mean error is 0.038 for synchronous, 0.046 for VLASH, and 0.071 for naive async.")
figure(os.path.join(SIMR, "fig_success.png"),
       "Fig. 5. Success rate versus inference delay. VLASH closely tracks the synchronous upper bound while the naive baseline collapses.")
figure(os.path.join(SIMR, "fig_tracking_error.png"),
       "Fig. 6. Mean tracking error versus inference delay.")
h3("3.3.3 Motion smoothness")
para("Figure 7 reports the action jerk. Naive asynchronous inference becomes progressively less smooth as the delay grows (jerk rising from 1.02 to 1.36 ×10⁻³), reflecting the corrective jumps that follow each over-shot chunk, whereas VLASH remains as smooth as synchronous inference (about 1.02–1.06 ×10⁻³) at all delays. This matches the qualitative behaviour in Figure 8: after each target jump the naive trace overshoots the new target position before correcting, whereas the VLASH and synchronous traces converge onto it cleanly.")
figure(os.path.join(SIMR, "fig_smoothness.png"),
       "Fig. 7. Motion smoothness (action jerk) versus inference delay; lower is smoother.")
figure(os.path.join(SIMR, "fig_trajectory.png"),
       "Fig. 8. End-effector tracking over time at Δ = 3 (x-coordinate). After each target jump the naive trace overshoots the new target before correcting, while VLASH and synchronous inference converge onto it cleanly.", width_cm=15.0)
h3("3.3.4 A residual gap to synchronous, explained by stale perception")
para("A small gap remains between VLASH and synchronous inference (69.5% versus 73.8% success at Δ = 4). This is expected and informative: the rollforward corrects the robot state exactly, but the target observation used by every asynchronous method is still the one captured when inference began, so it is stale by Δ steps. The testbed therefore reproduces, at small scale, precisely the limitation the VLASH authors identify — that future-state awareness compensates for proprioceptive change but not for changes in the visual scene.")

h2("3.4 Corroboration with Large-Scale VLASH Results")
para("The large-scale evaluation reported by the VLASH authors [1] confirms the same trends on a real VLA model and real robots. Table 3 reproduces the published π0.5 results across the four LIBERO sub-benchmarks. VLASH matches synchronous accuracy at low delay (97.2% at Δ = 1 and 97.1% at Δ = 2 versus the 96.8% synchronous baseline) while delivering 1.17× and 1.31× speed-ups, and degrades only gracefully at higher delay (94.6% at Δ = 3, 93.1% at Δ = 4) with speed-ups up to 1.47×. A notable secondary finding concerns the robot state: a model fine-tuned and run without state input actually scores slightly higher than the state-conditioned model under plain synchronous inference (97.7% versus 96.8%). This indicates that the base VLA under-utilises proprioceptive state, and it is exactly this observation that motivates the temporal-offset augmentation of Section 2.2.2 — the augmentation forces the model to attend to the state input so that, at deployment, it can exploit the rolled-forward state rather than ignore it.")
table_caption("Tab. 3. π0.5 on LIBERO under different inference delays (reported by [1]).")
make_table([
    ["Method", "Δ", "Spatial", "Object", "Goal", "LIBERO-10", "Avg SR (%)", "Speedup"],
    ["Sync", "0", "97.3", "99.6", "96.7", "93.5", "96.8", "1.00×"],
    ["Sync (w/o state)", "–", "98.5", "99.6", "97.3", "95.4", "97.7", "1.00×"],
    ["VLASH", "1", "98.8", "99.2", "96.7", "94.4", "97.2", "1.17×"],
    ["VLASH", "2", "97.5", "99.2", "97.3", "94.6", "97.1", "1.31×"],
    ["VLASH", "3", "94.4", "98.8", "93.3", "91.9", "94.6", "1.47×"],
    ["VLASH", "4", "92.5", "96.9", "93.3", "89.6", "93.1", "1.45×"],
], source="reported by [1]", font_size=9)
para("On the fast-physics Kinetix benchmark, designed to stress reaction under delay, the contrast is far sharper than on LIBERO. At an inference delay of four steps VLASH attains 81.7% success against only 51.2% for naive asynchronous inference — a 30.5 percentage-point improvement — and also surpasses Real-Time Chunking, which incurs additional inpainting overhead. This is the setting where the misalignment problem is most damaging, directly paralleling the steep collapse of the naive baseline in the simulation.")
para("On three physical tasks (pick-and-place, stacking, sorting) executed on real arms, VLASH attains the highest average score (94%), ahead of naive asynchronous inference (89.7%) and synchronous inference (83%), while completing tasks about 1.12× faster than synchronous control; synchronous inference scores lowest because its action stalls slow task execution under the time-aware scoring used. Adding action quantization raises the speed-up further: q = 2 reaches 2.03× with no accuracy loss, and q = 3 reaches 2.67× at the cost of a modest 4.7-point score reduction. Finally, because asynchronous inference overlaps computation with execution, it collapses the maximum reaction latency from “finish the current chunk, then infer” to “infer only”: on an RTX 4090 it drops from 536 ms to 36 ms (a 14.9× reduction), on a faster RTX 5090 the reduction reaches 17.4×, and even on a slower RTX 5070 it is 8.8×. These gains allow VLASH to support genuinely dynamic tasks such as a human–robot ping-pong rally, which are infeasible under synchronous control.")

h2("3.5 Discussion")
para("Why VLASH works. Both the simulation and the large-scale results point to the same explanation. Asynchronous inference is fast because it overlaps computation with execution, but that overlap is exactly what creates the misalignment: the chunk is committed against a state the robot has already left. VLASH removes this misalignment at its source, before the policy runs, by conditioning on the analytically rolled-forward state. Unlike methods that correct the chunk after the fact — RTC through runtime inpainting, or A2C2 through an added per-step correction head — VLASH adds no runtime computation beyond a vector summation and changes nothing in the model architecture. The simulation isolates this mechanism cleanly: the misalignment ε is driven to zero (Figure 4) and downstream accuracy follows (Figures 5–6).")
para("Agreement between small- and large-scale evidence. The controlled simulation and the full-scale evaluation agree on every qualitative claim: misalignment grows with delay; naive async collapses as delay increases, most severely in dynamic settings; VLASH tracks the synchronous upper bound; and a residual gap remains that is attributable to stale perception rather than to robot-state error. The convergence of an independent small-scale reconstruction with the published numbers strengthens confidence that the behaviour stems from the misalignment mechanism rather than from incidental properties of any one benchmark.")
para("Limitations. Three limitations should be stated plainly. First, the simulation is a kinematic abstraction: it has no contact dynamics, no real perception, and an idealised planner, so it validates the mechanism rather than predicting absolute success rates. Second, both the simulation and VLASH itself compensate only for robot state; neither predicts the future visual scene, so a residual gap to synchronous accuracy remains whenever the environment changes during inference. Third, the exactness of the rollforward depends on the delta-action representation; a Cartesian or absolute-joint action space would require forward kinematics and would not be free.")

h2("3.6 Summary of the Chapter")
para("This chapter presented two mutually reinforcing bodies of evidence. The original simulation testbed showed, under controlled and reproducible conditions, that prediction–execution misalignment grows with inference delay for naive asynchronous inference, that this misalignment is eliminated by state rollforward, and that the resulting accuracy and smoothness of VLASH track the synchronous upper bound while the naive baseline collapses (success at Δ = 4: 73.8% synchronous, 69.5% VLASH, 35.4% naive). The large-scale evaluation corroborated these trends on π0.5: comparable-to-synchronous accuracy with 1.17–1.47× speed-ups on LIBERO, a 30.5-point advantage over naive async on the dynamic Kinetix benchmark, the highest score on three real-world tasks, and reaction-latency reductions of up to 17.4×. A small residual gap, observed in both settings, was attributed to stale visual observation — the principal remaining limitation and a direction for future work.")

# ===================================================================
# CHAPTER 4
# ===================================================================
h1("CHAPTER 4: Conclusion and Future Work")
h2("4.0 Conclusion")
para("This thesis addressed a fundamental bottleneck in deploying large Vision-Language-Action models for real-time robotic control: the prediction–execution misalignment that arises when asynchronous inference is used to eliminate the action stall. It demonstrated that this misalignment — which causes naive asynchronous methods to produce jittery, inaccurate behaviour — can be resolved through a conceptually simple and computationally free technique, conditioning the policy on a future robot state computed analytically from the pending action queue (state rollforward).")
para("Modern VLAs such as π0.5 require tens to hundreds of milliseconds per inference, creating a delay of one to four control steps between when an observation is captured and when the resulting actions begin executing. Synchronous inference avoids misalignment but stalls the robot during every query; naive asynchronous inference removes the stall but leaves misalignment unaddressed; and prior corrective methods (RTC, A2C2) either add runtime overhead or add a trained correction module. The VLASH method resolves these limitations through state rollforward, temporal-offset augmentation, and a shared-observation training optimisation, adding nothing at deployment beyond a single vector summation.")
para("Assessment of the objective. The objectives set out in the Introduction were achieved. The thesis presented an overview of VLA models and real-time inference, analysed the misalignment problem, explained the future-state-aware method, and — as its practical contribution — designed and implemented an original simulation testbed that compares the three inference mechanisms. The testbed produced quantified, reproducible evidence: misalignment for naive asynchronous inference grows with delay (to ε = 0.036 at Δ = 4) and is eliminated by rollforward (ε = 0), while VLASH retains 69.5% success against the synchronous 73.8% and the collapsed 35.4% of naive async at Δ = 4. These results reproduce, at small scale, the trends reported for VLASH at full scale (comparable-to-synchronous accuracy with 1.17–1.47× speed-ups on LIBERO, +30.5 points over naive async on Kinetix, and reaction-latency reductions up to 17.4×), so the degree to which the objective was met can be considered high.")
para("Comparison with existing solutions. Relative to naive asynchronous inference, the studied approach restores accuracy without sacrificing the latency advantage. Relative to RTC it avoids runtime inpainting overhead, and relative to A2C2 it requires no added correction module and no architectural change. The principal weakness of the author’s own contribution is that the simulation is a kinematic abstraction without contact dynamics or real perception; it validates the mechanism rather than predicting absolute task performance, and its residual VLASH-versus-synchronous gap — caused by stale perception — mirrors the same limitation present in the full method.")
h2("4.1 Future Work")
para("Several directions remain open. The most significant limitation is that state rollforward compensates only for robot proprioception, not for changes in the visual scene; a natural extension is to predict the future observation with a learned world model or a lightweight video-prediction network, conditioning the policy on both a rolled-forward state and a predicted future image. A second direction is to validate the method beyond π0.5 and SmolVLA on newer architectures such as GR00T N1 [15], Gemini Robotics [17], and OpenVLA [16]. A third is edge deployment: combining VLASH with quantization-aware training and distillation to keep large policies misalignment-aware under the extreme latency of embedded hardware — including the author’s own RTX 3050 (6 GB) target, on which a real QLoRA fine-tune and LIBERO evaluation is the immediate next step. Adaptive inference scheduling and a learned, phase-dependent action-quantization policy are further refinements. Finally, pairing reactive inference with embodied memory such as Multi-Scale Embodied Memory [18] is a promising route toward robots that are both fast-reacting and capable of coherent long-horizon behaviour.")

# ===================================================================
# BIBLIOGRAPHY
# ===================================================================
h1("Bibliography")
REFS = [
 'J. Tang, Y. Sun, Y. Zhao, S. Yang, Y. Lin, Z. Zhang, J. Hou, Y. Lu, Z. Liu, and S. Han, "VLASH: Real-Time VLAs via Future-State-Aware Asynchronous Inference," arXiv:2512.01031, 2025.',
 'K. Black, N. Brown, D. Driess, A. Esmail, M. Equi, C. Finn, et al., "π0: A Vision-Language-Action Flow Model for General Robot Control," arXiv:2410.24164, 2024.',
 'Physical Intelligence, K. Black, et al., "π0.5: a Vision-Language-Action Model with Open-World Generalization," arXiv:2504.16054, 2025.',
 'M. Shukor, D. Aubakirova, F. Capuano, P. Kooijmans, S. Palma, A. Zouitine, et al., "SmolVLA: A Vision-Language-Action Model for Affordable and Efficient Robotics," arXiv:2506.01844, 2025.',
 'K. Black, M. Y. Galliker, and S. Levine, "Real-Time Execution of Action Chunking Flow Policies," arXiv:2506.07339; NeurIPS, 2025.',
 '"Leave No Observation Behind: Real-Time Correction for VLA Action Chunks (A2C2)," arXiv:2509.23224, 2025.',
 'B. Liu, Y. Zhu, C. Gao, Y. Feng, Q. Liu, Y. Zhu, and P. Stone, "LIBERO: Benchmarking Knowledge Transfer for Lifelong Robot Learning," in Advances in Neural Information Processing Systems (NeurIPS), vol. 36, 2023.',
 'T. Z. Zhao, V. Kumar, S. Levine, and C. Finn, "Learning Fine-Grained Bimanual Manipulation with Low-Cost Hardware (ACT)," in Robotics: Science and Systems (RSS), 2023.',
 'M. Matthews, M. Beukman, F. Christianos, L. Schäfer, M. Foster, et al., "Kinetix: Investigating the Training of General Agents through Open-Ended Physics-Based Control Tasks," in International Conference on Learning Representations (ICLR), 2025.',
 'E. J. Hu, Y. Shen, P. Wallis, Z. Allen-Zhu, Y. Li, S. Wang, L. Wang, and W. Chen, "LoRA: Low-Rank Adaptation of Large Language Models," in International Conference on Learning Representations (ICLR), 2022.',
 'Y. Lipman, R. T. Q. Chen, H. Ben-Hamu, M. Nickel, and M. Le, "Flow Matching for Generative Modeling," in International Conference on Learning Representations (ICLR), 2023.',
 'A. Brohan, N. Brown, J. Carbajal, Y. Chebotar, et al., "RT-2: Vision-Language-Action Models Transfer Web Knowledge to Robotic Control," in Conference on Robot Learning (CoRL), 2023.',
 'L. Beyer, A. Steiner, A. S. Pinto, A. Kolesnikov, X. Wang, et al., "PaliGemma: A Versatile 3B VLM for Transfer," arXiv:2407.07519, 2024.',
 'X. Zhai, B. Mustafa, A. Kolesnikov, and L. Beyer, "Sigmoid Loss for Language Image Pre-Training (SigLIP)," in IEEE/CVF International Conference on Computer Vision (ICCV), 2023.',
 'NVIDIA, "GR00T N1: An Open Foundation Model for Generalist Humanoid Robots," arXiv:2503.14734, 2025.',
 'M. J. Kim, K. Pertsch, S. Karamcheti, et al., "OpenVLA: An Open-Source Vision-Language-Action Model," arXiv:2406.09246, 2024.',
 'A. Abdolmaleki, S. Abeyruwan, J. Ainslie, et al., "Gemini Robotics 1.5: Pushing the Frontier of Generalist Robots," arXiv:2510.03342, 2025.',
 'M. Torne, K. Pertsch, H. Walke, K. Vedder, S. Nair, B. Ichter, et al., "Multi-Scale Embodied Memory (MEM): VLAs with Long and Short-Term Memory," Physical Intelligence, 2026.',
 'R. Cadene, S. Alibert, A. Soare, Q. Gallouedec, et al., "LeRobot: State-of-the-Art Machine Learning for Real-World Robotics in PyTorch," github.com/huggingface/lerobot, 2024.',
]
for i, ref in enumerate(REFS, 1):
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.8); p.paragraph_format.first_line_indent = Cm(-0.8)
    r = p.add_run(f"[{i}] {ref}"); r.font.name = FONT; r.font.size = Pt(12)

# ===================================================================
# LIST OF FIGURES / TABLES (static)
# ===================================================================
h1("List of Figures")
FIGLIST = [
 "Fig. 1. Architecture of the π0.5 vision-language-action model.",
 "Fig. 2. Three inference paradigms: synchronous, naive asynchronous, and VLASH.",
 "Fig. 3. State rollforward — computing the execution-time state by vector summation.",
 "Fig. 4. Prediction–execution misalignment ε versus inference delay.",
 "Fig. 5. Success rate versus inference delay.",
 "Fig. 6. Mean tracking error versus inference delay.",
 "Fig. 7. Motion smoothness (action jerk) versus inference delay.",
 "Fig. 8. End-effector tracking over time at Δ = 3 (x-coordinate).",
]
for f in FIGLIST:
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.15; p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f); r.font.name = FONT; r.font.size = Pt(12)

h1("List of Tables")
TABLIST = [
 "Tab. 1. Measured inference latency of π0.5 on representative hardware.",
 "Tab. 2. Simulation results (mean over 50 trials per delay).",
 "Tab. 3. π0.5 on LIBERO under different inference delays (reported by the VLASH authors).",
]
for tline in TABLIST:
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.15; p.paragraph_format.space_after = Pt(3)
    r = p.add_run(tline); r.font.name = FONT; r.font.size = Pt(12)

# ===================================================================
# SUMMARY (abstract, at the very end)
# ===================================================================
h1("Summary")
para("This thesis studies real-time inference for Vision-Language-Action (VLA) models in reactive robotics. Action-chunking VLAs such as π0.5 are accurate but slow to deploy: synchronous inference stalls the robot during every model query, while naive asynchronous inference removes the stall at the cost of a prediction–execution misalignment, because the policy conditions on a state the robot has already left by the time its actions execute. The thesis analyses this misalignment and studies the VLASH method, which makes the policy future-state-aware by rolling the robot state forward through the already-known pending actions — an exact, zero-overhead vector summation enabled by the delta-action representation.")
para("As its practical contribution, the thesis designs and implements an original, fully reproducible simulation testbed that compares synchronous, naive asynchronous, and future-state-aware inference on a reactive reaching task while sweeping the inference delay. The experiments show that misalignment grows with delay for naive asynchronous inference and is eliminated by state rollforward, that VLASH retains task accuracy and motion smoothness close to the synchronous upper bound while the naive baseline collapses, and that a small residual gap is attributable to stale visual perception. These findings reproduce, at small scale, the trends reported for VLASH at full scale, and the thesis concludes with directions for future work, including visual-observation rollforward and deployment on resource-constrained hardware.")
para("Keywords: Vision-Language-Action models, asynchronous inference, real-time robotics, state rollforward, action chunking, π0.5, VLASH.", indent=False)

doc.save(OUTDOC)
print("FULL THESIS saved:", OUTDOC)
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))
